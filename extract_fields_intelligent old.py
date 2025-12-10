#!/usr/bin/env python3
"""
Document Field Extraction System using GPT API
Complete single-file implementation with intelligent column selection

This script passes the COMPLETE configuration to the LLM, which then selects
the appropriate column based on document name similarity. Makes ONE LLM call 
per document and merges all results into a single Excel output.

Key Feature: LLM intelligently matches document name to config column

Requirements:
    pip install pandas openpyxl python-docx pypdf pdfplumber openai

Usage:
    export OPENAI_API_KEY='your-api-key-here'
    python extract_fields.py
"""

import os
import json
import sys
from typing import Dict, List, Optional, Tuple
from pathlib import Path
from datetime import datetime
from dotenv import load_dotenv

# Document processing imports
import pandas as pd
from openpyxl import load_workbook
from docx import Document
from pypdf import PdfReader
import pdfplumber
from openai import OpenAI

load_dotenv()


# =============================================================================
# CONFIGURATION SECTION - MODIFY THESE PATHS
# =============================================================================

CONFIG_FILE_PATH = "FieldConfigrationFile.xlsx"  # Configuration Excel file
DOCUMENTS_FOLDER = "email_attachments/3000023987/20251207_173848"                   # Folder containing documents to process
OUTPUT_FILE_PATH = "extraction_results.xlsx"  # Output Excel file
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")            # OpenAI API key
OPENAI_MODEL = "gpt-4o"                                      # Model: gpt-4o-mini or gpt-4o (use gpt-4o for better column selection)
MAX_DOCUMENT_CHARS = 15000                                   # Max characters to send to GPT

print("OPENAI_API_KEY: ",OPENAI_API_KEY)
# =============================================================================
# DOCUMENT READING FUNCTIONS
# =============================================================================

def read_pdf_document(file_path: str) -> str:
    """Extract text from PDF document including tables"""
    text_parts = []
    
    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"\n{'='*60}\n[PAGE {page_num}]\n{'='*60}\n")
                    text_parts.append(page_text)
                
                tables = page.extract_tables()
                if tables:
                    for table_num, table in enumerate(tables, 1):
                        text_parts.append(f"\n[TABLE {table_num} ON PAGE {page_num}]\n")
                        for row in table:
                            if row:
                                row_text = " | ".join(str(cell) if cell else "" for cell in row)
                                text_parts.append(row_text)
                        text_parts.append("")
        
        return "\n".join(text_parts)
    
    except Exception as e:
        print(f"  ⚠ Error reading PDF {file_path}: {e}")
        return ""


def read_word_document(file_path: str) -> str:
    """Extract text from Word document including tables"""
    text_parts = []
    
    try:
        doc = Document(file_path)
        
        text_parts.append("="*60)
        text_parts.append("[DOCUMENT CONTENT]")
        text_parts.append("="*60)
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        if doc.tables:
            text_parts.append("\n" + "="*60)
            text_parts.append("[TABLES]")
            text_parts.append("="*60)
            
            for table_num, table in enumerate(doc.tables, 1):
                text_parts.append(f"\n[TABLE {table_num}]")
                for row_num, row in enumerate(table.rows):
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    text_parts.append(row_text)
        
        return "\n".join(text_parts)
    
    except Exception as e:
        print(f"  ⚠ Error reading Word document {file_path}: {e}")
        return ""


def read_excel_document(file_path: str) -> str:
    """Extract text from Excel document"""
    text_parts = []
    
    try:
        wb = load_workbook(file_path, data_only=True)
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            
            text_parts.append("="*60)
            text_parts.append(f"[SHEET: {sheet_name}]")
            text_parts.append("="*60)
            
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                if row_text.strip(" |"):
                    text_parts.append(row_text)
        
        return "\n".join(text_parts)
    
    except Exception as e:
        print(f"  ⚠ Error reading Excel file {file_path}: {e}")
        return ""


def read_document(file_path: str) -> str:
    """Read document based on file extension"""
    ext = Path(file_path).suffix.lower()
    
    if ext == '.pdf':
        return read_pdf_document(file_path)
    elif ext in ['.docx', '.doc']:
        return read_word_document(file_path)
    elif ext in ['.xlsx', '.xls']:
        return read_excel_document(file_path)
    else:
        print(f"  ⚠ Unsupported file format: {ext}")
        return ""


# =============================================================================
# CONFIGURATION MANAGEMENT
# =============================================================================

def load_configuration(config_path: str) -> Tuple[pd.DataFrame, List[str]]:
    """
    Load configuration file and extract PAS field names
    
    Returns:
        Tuple of (config_dataframe, list_of_pas_fields)
    """
    try:
        config_df = pd.read_excel(config_path, sheet_name='Sheet1')
        pas_fields = config_df['PAS Field Name'].tolist()
        return config_df, pas_fields
    except Exception as e:
        print(f"✗ Error loading configuration file: {e}")
        sys.exit(1)


def prepare_config_for_llm(config_df: pd.DataFrame) -> str:
    """
    Prepare configuration data in a format suitable for LLM
    
    Returns:
        Formatted string representation of configuration
    """
    config_text = []
    
    # Get description columns (exclude PAS Field Name, Data Type, Field length)
    description_columns = [col for col in config_df.columns 
                          if 'Description' in col]
    
    config_text.append("CONFIGURATION FILE STRUCTURE:")
    config_text.append("="*80)
    config_text.append(f"\nAvailable instruction columns: {len(description_columns)}")
    for i, col in enumerate(description_columns, 1):
        # Count how many fields have descriptions in this column
        field_count = config_df[col].notna().sum()
        config_text.append(f"{i}. {col} ({field_count} fields with instructions)")
    
    config_text.append("\n" + "="*80)
    config_text.append("FIELD EXTRACTION INSTRUCTIONS:")
    config_text.append("="*80)
    
    # For each PAS field, show all available descriptions
    for idx, row in config_df.iterrows():
        field_name = row['PAS Field Name']
        config_text.append(f"\n[FIELD: {field_name}]")
        
        # Show description from each column
        for col in description_columns:
            description = row.get(col, '')
            if pd.notna(description) and str(description).strip():
                config_text.append(f"  • {col}: {str(description)}")
    
    return "\n".join(config_text)


# =============================================================================
# GPT EXTRACTION - ONE CALL PER DOCUMENT WITH FULL CONFIG
# =============================================================================

def extract_fields_with_intelligent_selection(document_text: str, 
                                              config_structure: str,
                                              document_name: str, 
                                              file_extension: str,
                                              pas_fields: List[str],
                                              api_key: str, 
                                              model: str) -> Tuple[Dict[str, str], str]:
    """
    Extract all fields from document using a SINGLE GPT API call
    LLM intelligently selects the appropriate configuration column
    
    Args:
        document_text: Full text content of the document
        config_structure: Complete configuration with all columns
        document_name: Name of the document (without extension)
        file_extension: File extension (.pdf, .docx, etc.)
        pas_fields: List of all PAS field names
        api_key: OpenAI API key
        model: Model to use (gpt-4o or gpt-4o-mini)
        
    Returns:
        Tuple of (extracted_fields_dict, selected_column_name)
    """
    
    # Truncate document if too long
    if len(document_text) > MAX_DOCUMENT_CHARS:
        print(f"  ℹ Document truncated from {len(document_text)} to {MAX_DOCUMENT_CHARS} characters")
        document_text = document_text[:MAX_DOCUMENT_CHARS]
    
    # Build the extraction prompt with intelligent column selection
    prompt = f"""You are an expert document field extraction system with intelligent configuration selection.

TASK OVERVIEW:
You will receive a document and a complete configuration file with multiple instruction columns. Your job is to:
1. Analyze the document name and type to select the MOST APPROPRIATE instruction column
2. Extract all fields using the instructions from that selected column
3. Return the extracted data along with the column you selected

DOCUMENT INFORMATION:
- Document Name: {document_name}
- File Extension: {file_extension}
- Document Type: {"Word Document" if file_extension in ['.docx', '.doc'] else "PDF Document" if file_extension == '.pdf' else "Excel Spreadsheet"}

COLUMN SELECTION RULES (IMPORTANT):
1. Match the document name to the closest instruction column name
2. SPECIAL CASE FOR PD DOCUMENTS:
   - If document name contains "PD" AND file is NOT a Word document (.docx/.doc) → Use "PD Description"
   - If document name contains "PD" AND file IS a Word document (.docx/.doc) → Use "PD (Word Doc) Description"
3. For other documents, find the column whose name is most similar to the document name:
   - "CAM" in document name → Use "CAM Description"
   - "Application" in document name → Use "Application Form Description"
   - "Legal" in document name → Use "Legal Doc Description"
   - "Technical" in document name → Use "Technical Doc Description"
   - "Email" in document name → Use "Email Subject Description" or "Email Body Description"
4. If no good match, use the column with the most field instructions available

{config_structure}

DOCUMENT CONTENT:
{document_text}

EXTRACTION INSTRUCTIONS:
1. FIRST: Analyze the document name "{document_name}" and type "{file_extension}"
2. SELECT the most appropriate instruction column based on the rules above
3. For each of the {len(pas_fields)} PAS fields, extract the value using instructions from your selected column
4. If a field has no instruction in the selected column, mark as "NO INSTRUCTION"
5. If a field has instruction but value is not found in document, mark as "NOT FOUND"
6. Extract exact values as they appear in the document
7. Do NOT make assumptions or infer values not explicitly stated

OUTPUT FORMAT:
Return a JSON object with TWO keys:
1. "selected_column": The name of the configuration column you selected
2. "extracted_fields": An object with field names as keys and extracted values as values

Example format:
{{
  "selected_column": "PD (Word Doc) Description",
  "extracted_fields": {{
    "Loan_Number": "LA-2024-12345",
    "Sourcing_Region": "North Region",
    "Branch_Name": "Main Branch",
    "Field_Without_Instruction": "NO INSTRUCTION",
    "Field_Not_In_Document": "NOT FOUND"
  }}
}}

CRITICAL: Include ALL {len(pas_fields)} PAS fields in your response, even if they have "NO INSTRUCTION" or are "NOT FOUND".

Return ONLY the JSON object, no additional text."""

    try:
        client = OpenAI(api_key=api_key)
        
        print(f"  → Making GPT API call with full configuration...")
        print(f"  → LLM will intelligently select the appropriate column...")
        
        response = client.chat.completions.create(
            model=model,
            messages=[
                {
                    "role": "system", 
                    "content": "You are a precise document field extraction assistant with intelligent configuration selection. Always respond with valid JSON."
                },
                {
                    "role": "user", 
                    "content": prompt
                }
            ],
            temperature=0.1,
            response_format={"type": "json_object"}
        )
        
        result_text = response.choices[0].message.content
        result_json = json.loads(result_text)
        print("Lets see the result_json:", result_json)
        
        selected_column = result_json.get("selected_column", "UNKNOWN")
        extracted_fields = result_json.get("extracted_fields", {})
        
        # Ensure all PAS fields are present
        for field in pas_fields:
            if field not in extracted_fields:
                extracted_fields[field] = "ERROR"
        
        print(f"  ✓ LLM selected column: '{selected_column}'")
        print(f"  ✓ Successfully extracted {len(extracted_fields)} fields")
        
        return extracted_fields, selected_column
    
    except json.JSONDecodeError as e:
        print(f"  ✗ JSON parsing error: {e}")
        return {field: "ERROR" for field in pas_fields}, "ERROR"
    
    except Exception as e:
        print(f"  ✗ API error: {e}")
        return {field: "ERROR" for field in pas_fields}, "ERROR"


# =============================================================================
# RESULT MERGING AND OUTPUT
# =============================================================================

def merge_results_to_excel(all_results: Dict[str, Dict[str, str]], 
                          pas_fields: List[str], 
                          output_path: str,
                          column_selections: Dict[str, str]) -> None:
    """
    Merge all extraction results into a single Excel file
    
    Args:
        all_results: Dictionary mapping document names to their extracted fields
        pas_fields: List of all PAS field names
        output_path: Path to save output Excel file
        column_selections: Dictionary mapping document names to selected columns
    """
    results_data = []
    
    for field in pas_fields:
        row = {'PAS Field Name': field}
        for doc_name, extracted_data in all_results.items():
            row[doc_name] = extracted_data.get(field, "NOT PROCESSED")
        results_data.append(row)
    
    results_df = pd.DataFrame(results_data)
    
    # Create metadata sheet with column selections
    metadata_data = []
    for doc_name, selected_col in column_selections.items():
        metadata_data.append({
            'Document Name': doc_name,
            'Selected Configuration Column': selected_col
        })
    metadata_df = pd.DataFrame(metadata_data)
    
    # Save to Excel with multiple sheets
    try:
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            results_df.to_excel(writer, sheet_name='Extracted Fields', index=False)
            metadata_df.to_excel(writer, sheet_name='Column Selections', index=False)
        
        print(f"\n✓ Results saved to: {output_path}")
        print(f"  - Sheet 1: Extracted Fields")
        print(f"  - Sheet 2: Column Selections (shows which config column was used)")
    except Exception as e:
        print(f"\n✗ Error saving results: {e}")


def display_summary(all_results: Dict[str, Dict[str, str]], 
                   pas_fields: List[str],
                   column_selections: Dict[str, str]) -> None:
    """Display summary of extraction results"""
    print("\n" + "="*80)
    print("EXTRACTION SUMMARY")
    print("="*80)
    print(f"Total PAS Fields: {len(pas_fields)}")
    print(f"Total Documents Processed: {len(all_results)}")
    
    print(f"\nColumn Selections by LLM:")
    for i, (doc_name, selected_col) in enumerate(column_selections.items(), 1):
        print(f"  {i}. {doc_name:30s} → {selected_col}")
    
    # Count extraction statistics
    total_extracted = 0
    total_not_found = 0
    total_no_instruction = 0
    total_errors = 0
    
    for doc_name, extracted_data in all_results.items():
        for value in extracted_data.values():
            if value == "NOT FOUND":
                total_not_found += 1
            elif value == "NO INSTRUCTION":
                total_no_instruction += 1
            elif value == "ERROR":
                total_errors += 1
            elif value != "NOT PROCESSED":
                total_extracted += 1
    
    total_cells = len(pas_fields) * len(all_results)
    
    print(f"\nExtraction Statistics:")
    print(f"  Fields Extracted: {total_extracted} ({total_extracted/total_cells*100:.1f}%)")
    print(f"  Not Found: {total_not_found} ({total_not_found/total_cells*100:.1f}%)")
    print(f"  No Instruction: {total_no_instruction} ({total_no_instruction/total_cells*100:.1f}%)")
    print(f"  Errors: {total_errors} ({total_errors/total_cells*100:.1f}%)")
    print("="*80)


# =============================================================================
# MAIN EXECUTION
# =============================================================================

def main():
    """Main execution function"""
    print("\n" + "="*80)
    print("DOCUMENT FIELD EXTRACTION SYSTEM")
    print("Intelligent Column Selection by LLM")
    print("="*80)
    print(f"\nTimestamp: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    # Validate API key
    if not OPENAI_API_KEY or OPENAI_API_KEY == "":
        print("\n✗ ERROR: OpenAI API key not set!")
        print("  Please set the OPENAI_API_KEY environment variable:")
        print("  export OPENAI_API_KEY='your-api-key-here'")
        sys.exit(1)
    
    # Load configuration
    print(f"\n[1/5] Loading Configuration")
    print(f"  Configuration file: {CONFIG_FILE_PATH}")
    
    if not os.path.exists(CONFIG_FILE_PATH):
        print(f"  ✗ Configuration file not found: {CONFIG_FILE_PATH}")
        sys.exit(1)
    
    config_df, pas_fields = load_configuration(CONFIG_FILE_PATH)
    print(f"  ✓ Loaded {len(pas_fields)} PAS fields")
    print(f"  ✓ Available columns: {len(config_df.columns)}")
    
    # Prepare configuration for LLM
    print(f"\n  Preparing complete configuration for LLM...")
    config_structure = prepare_config_for_llm(config_df)
    print(f"  ✓ Configuration prepared ({len(config_structure)} characters)")
    
    # Check documents folder
    print(f"\n[2/5] Scanning Documents")
    print(f"  Documents folder: {DOCUMENTS_FOLDER}")
    
    if not os.path.exists(DOCUMENTS_FOLDER):
        print(f"  ✗ Documents folder not found!")
        print(f"  Creating folder: {DOCUMENTS_FOLDER}")
        os.makedirs(DOCUMENTS_FOLDER)
        print(f"  Please place your documents in this folder and run again.")
        sys.exit(1)
    
    supported_extensions = ['.pdf', '.docx', '.doc', '.xlsx', '.xls']
    documents = [
        f for f in os.listdir(DOCUMENTS_FOLDER)
        if Path(f).suffix.lower() in supported_extensions
    ]
    
    if not documents:
        print(f"  ✗ No documents found in {DOCUMENTS_FOLDER}")
        print(f"  Supported formats: {', '.join(supported_extensions)}")
        sys.exit(1)
    
    print(f"  ✓ Found {len(documents)} document(s) to process:")
    for i, doc in enumerate(documents, 1):
        print(f"    {i}. {doc}")
    
    # Process each document
    print(f"\n[3/5] Processing Documents")
    print(f"  Using model: {OPENAI_MODEL}")
    print(f"  Note: LLM will select appropriate config column for each document")
    
    all_results = {}
    column_selections = {}
    
    for doc_idx, doc_file in enumerate(documents, 1):
        print(f"\n{'='*80}")
        print(f"DOCUMENT {doc_idx}/{len(documents)}: {doc_file}")
        print(f"{'='*80}")
        
        doc_path = os.path.join(DOCUMENTS_FOLDER, doc_file)
        doc_name = Path(doc_file).stem
        doc_ext = Path(doc_file).suffix
        
        # Step 1: Read document
        print(f"\n[Step 1] Reading Document Content")
        document_text = read_document(doc_path)
        
        if not document_text or len(document_text) < 50:
            print(f"  ✗ Could not extract meaningful content from document")
            print(f"  Skipping document...")
            continue
        
        print(f"  ✓ Extracted {len(document_text)} characters")
        print(f"  ✓ Preview: {document_text[:100]}...")
        
        # Step 2: Extract fields with intelligent column selection
        print(f"\n[Step 2] Extracting Fields with Intelligent Column Selection")
        print(f"  Document name: {doc_name}")
        print(f"  File type: {doc_ext}")
        print(f"  This will make ONE API call that:")
        print(f"    1. Analyzes document name and type")
        print(f"    2. Selects appropriate configuration column")
        print(f"    3. Extracts all {len(pas_fields)} fields")
        
        extracted_data, selected_column = extract_fields_with_intelligent_selection(
            document_text=document_text,
            config_structure=config_structure,
            document_name=doc_name,
            file_extension=doc_ext,
            pas_fields=pas_fields,
            api_key=OPENAI_API_KEY,
            model=OPENAI_MODEL
        )
        
        # Step 3: Store results
        if extracted_data:
            all_results[doc_name] = extracted_data
            column_selections[doc_name] = selected_column
            print(f"\n  ✓ Completed extraction for '{doc_file}'")
            
            # Show sample of extracted data
            sample_count = min(5, len(extracted_data))
            print(f"\n  Sample extracted fields ({sample_count}/{len(extracted_data)}):")
            sample_items = list(extracted_data.items())[:sample_count]
            for i, (field, value) in enumerate(sample_items, 1):
                display_value = value if len(str(value)) <= 50 else str(value)[:47] + "..."
                print(f"    {i}. {field}: {display_value}")
    
    # Merge results
    if not all_results:
        print("\n" + "="*80)
        print("✗ No results to save - no documents were successfully processed")
        print("="*80)
        sys.exit(1)
    
    print(f"\n[4/5] Merging Results")
    print(f"  Total documents processed: {len(all_results)}")
    print(f"  Total PAS fields: {len(pas_fields)}")
    
    # Create output directory if it doesn't exist
    output_dir = Path(OUTPUT_FILE_PATH).parent
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    merge_results_to_excel(all_results, pas_fields, OUTPUT_FILE_PATH, column_selections)
    
    # Display summary
    print(f"\n[5/5] Generating Summary")
    display_summary(all_results, pas_fields, column_selections)
    
    # Final output
    print(f"\n{'='*80}")
    print("EXTRACTION COMPLETE!")
    print(f"{'='*80}")
    print(f"\nOutput file: {OUTPUT_FILE_PATH}")
    print(f"Total documents: {len(all_results)}")
    print(f"Total API calls made: {len(all_results)} (one per document)")
    print(f"\nThe Excel file contains:")
    print(f"  • Sheet 1: Extracted Fields (165 fields × {len(all_results)} documents)")
    print(f"  • Sheet 2: Column Selections (which config column was used for each document)")
    print(f"{'='*80}\n")


# =============================================================================
# ENTRY POINT
# =============================================================================

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n\n✗ Process interrupted by user")
        sys.exit(1)
    except Exception as e:
        print(f"\n✗ Unexpected error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
