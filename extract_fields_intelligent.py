#!/usr/bin/env python3
"""
Document Field Extraction System using GPT API
Complete single-file implementation with LLM-based intelligent column selection

This script passes the COMPLETE configuration to the LLM, which then intelligently
analyzes the document and selects the appropriate column based on document name,
type, and content. Makes ONE LLM call per document and merges all results into 
Excel files saved in each respective folder.

Key Features:
- LLM intelligently matches document to appropriate config column
- Recursive folder processing - automatically discovers all document folders
- Saves results in each folder alongside source documents

Requirements:
    pip install pandas openpyxl python-docx pypdf pdfplumber openai

Usage:
    export OPENAI_API_KEY='your-api-key-here'
    python extract_fields_intelligent.py
"""

import os
import json
import sys
from typing import Dict, List, Optional, Tuple
from pathlib import Path
from datetime import datetime
from dotenv import load_dotenv

# Document processing imports
import pandas as pd
from openpyxl import load_workbook
from docx import Document
from pypdf import PdfReader
import pdfplumber
from openai import OpenAI

load_dotenv()


# =============================================================================
# CONFIGURATION SECTION - MODIFY THESE PATHS
# =============================================================================

CONFIG_FILE_PATH = "FieldConfigrationFile.xlsx"  # Configuration Excel file
BASE_DOCUMENTS_FOLDER = "email_attachments"      # Base folder - will process all subfolders recursively
OUTPUT_FILENAME = "extraction_results.xlsx"      # Output Excel filename (created in each folder)
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")            # OpenAI API key
OPENAI_MODEL = "gpt-4o"                                      # Model: gpt-4o-mini or gpt-4o (use gpt-4o for better accuracy)
MAX_DOCUMENT_CHARS = 15000                                   # Max characters to send to GPT


# =============================================================================
# DOCUMENT READING FUNCTIONS
# =============================================================================

def read_pdf_document(file_path: str) -> str:
    """Extract text from PDF document including tables"""
    text_parts = []
    
    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"\n{'='*60}\n[PAGE {page_num}]\n{'='*60}\n")
                    text_parts.append(page_text)
                
                tables = page.extract_tables()
                if tables:
                    for table_num, table in enumerate(tables, 1):
                        text_parts.append(f"\n[TABLE {table_num} ON PAGE {page_num}]\n")
                        for row in table:
                            if row:
                                row_text = " | ".join(str(cell) if cell else "" for cell in row)
                                text_parts.append(row_text)
                        text_parts.append("")
        
        return "\n".join(text_parts)
    
    except Exception as e:
        print(f"  ⚠ Error reading PDF {file_path}: {e}")
        return ""


def read_word_document(file_path: str) -> str:
    """Extract text from Word document including tables"""
    text_parts = []
    
    try:
        doc = Document(file_path)
        
        text_parts.append("="*60)
        text_parts.append("[DOCUMENT CONTENT]")
        text_parts.append("="*60)
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        if doc.tables:
            text_parts.append("\n" + "="*60)
            text_parts.append("[TABLES]")
            text_parts.append("="*60)
            
            for table_num, table in enumerate(doc.tables, 1):
                text_parts.append(f"\n[TABLE {table_num}]")
                for row_num, row in enumerate(table.rows):
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    text_parts.append(row_text)
        
        return "\n".join(text_parts)
    
    except Exception as e:
        print(f"  ⚠ Error reading Word document {file_path}: {e}")
        return ""


def read_excel_document(file_path: str) -> str:
    """Extract text from Excel document"""
    text_parts = []
    
    try:
        wb = load_workbook(file_path, data_only=True)
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            
            text_parts.append("="*60)
            text_parts.append(f"[SHEET: {sheet_name}]")
            text_parts.append("="*60)
            
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                if row_text.strip(" |"):
                    text_parts.append(row_text)
        
        return "\n".join(text_parts)
    
    except Exception as e:
        print(f"  ⚠ Error reading Excel file {file_path}: {e}")
        return ""


def read_document(file_path: str) -> str:
    """Read document based on file extension"""
    ext = Path(file_path).suffix.lower()
    
    if ext == '.pdf':
        return read_pdf_document(file_path)
    elif ext in ['.docx', '.doc']:
        return read_word_document(file_path)
    elif ext in ['.xlsx', '.xls']:
        return read_excel_document(file_path)
    else:
        print(f"  ⚠ Unsupported file format: {ext}")
        return ""


# =============================================================================
# FOLDER DISCOVERY AND PROCESSING
# =============================================================================

def find_document_folders(base_folder: str, supported_extensions: List[str]) -> List[str]:
    """
    Recursively find all folders containing supported documents
    
    Args:
        base_folder: Base directory to search
        supported_extensions: List of file extensions to look for
        
    Returns:
        List of folder paths containing documents
    """
    document_folders = []
    
    if not os.path.exists(base_folder):
        return document_folders
    
    # Walk through all directories
    for root, dirs, files in os.walk(base_folder):
        # Check if this folder has any supported documents
        has_documents = any(
            Path(f).suffix.lower() in supported_extensions 
            for f in files
        )
        
        if has_documents:
            document_folders.append(root)
    
    return document_folders


def get_documents_in_folder(folder_path: str, supported_extensions: List[str]) -> List[str]:
    """
    Get list of supported documents in a specific folder (non-recursive)
    
    Args:
        folder_path: Path to the folder
        supported_extensions: List of file extensions to look for
        
    Returns:
        List of document filenames (not full paths)
    """
    if not os.path.exists(folder_path):
        return []
    
    documents = [
        f for f in os.listdir(folder_path)
        if os.path.isfile(os.path.join(folder_path, f)) and
        Path(f).suffix.lower() in supported_extensions
    ]
    
    return sorted(documents)


def process_single_folder(folder_path: str,
                         config_structure: str,
                         pas_fields: List[str],
                         api_key: str,
                         model: str,
                         output_filename: str) -> Tuple[bool, str]:
    """
    Process all documents in a single folder and save results there
    
    Args:
        folder_path: Path to folder containing documents
        config_structure: Complete configuration formatted for LLM (all columns)
        pas_fields: List of PAS field names
        api_key: OpenAI API key
        model: Model to use
        output_filename: Name of output Excel file
        
    Returns:
        Tuple of (success: bool, message: str)
    """
    supported_extensions = ['.pdf', '.docx', '.doc', '.xlsx', '.xls']
    
    # Get documents in this folder only (not subfolders)
    documents = get_documents_in_folder(folder_path, supported_extensions)
    
    if not documents:
        return False, "No documents found"
    
    print(f"\n{'='*80}")
    print(f"PROCESSING FOLDER: {folder_path}")
    print(f"{'='*80}")
    print(f"Found {len(documents)} document(s):")
    for i, doc in enumerate(documents, 1):
        print(f"  {i}. {doc}")
    
    all_results = {}
    column_selections = {}
    
    # Process each document
    for doc_idx, doc_file in enumerate(documents, 1):
        print(f"\n{'-'*80}")
        print(f"Document {doc_idx}/{len(documents)}: {doc_file}")
        print(f"{'-'*80}")
        
        doc_path = os.path.join(folder_path, doc_file)
        doc_name = Path(doc_file).stem
        doc_ext = Path(doc_file).suffix
        
        # Read document
        print(f"  [1/2] Reading document...")
        document_text = read_document(doc_path)
        
        if not document_text or len(document_text) < 50:
            print(f"  ✗ Could not extract meaningful content")
            print(f"  Skipping document...")
            continue
        
        print(f"  ✓ Extracted {len(document_text)} characters")
        
        # Extract fields - LLM will select column
        print(f"  [2/2] Extracting {len(pas_fields)} fields with intelligent column selection...")
        extracted_data, selected_column = extract_fields_with_intelligent_selection(
            document_text=document_text,
            config_structure=config_structure,
            document_name=doc_name,
            file_extension=doc_ext,
            pas_fields=pas_fields,
            api_key=api_key,
            model=model
        )
        
        if extracted_data:
            all_results[doc_name] = extracted_data
            column_selections[doc_name] = selected_column
            print(f"  ✓ Extraction completed")
    
    # Save results in the same folder
    if all_results:
        output_path = os.path.join(folder_path, output_filename)
        merge_results_to_excel(all_results, pas_fields, output_path, column_selections)
        
        # Display summary
        print(f"\n{'='*80}")
        print(f"FOLDER SUMMARY: {folder_path}")
        print(f"{'='*80}")
        print(f"Documents processed: {len(all_results)}/{len(documents)}")
        print(f"Output saved to: {output_filename}")
        print(f"{'='*80}\n")
        
        return True, f"Processed {len(all_results)} documents"
    else:
        print(f"\n✗ No documents were successfully processed in this folder")
        return False, "No successful extractions"


# =============================================================================
# CONFIGURATION MANAGEMENT
# =============================================================================

def load_configuration(config_path: str) -> Tuple[pd.DataFrame, List[str]]:
    """
    Load configuration file and extract PAS field names
    
    Returns:
        Tuple of (config_dataframe, list_of_pas_fields)
    """
    try:
        config_df = pd.read_excel(config_path, sheet_name='Sheet1')
        pas_fields = config_df['PAS Field Name'].tolist()
        return config_df, pas_fields
    except Exception as e:
        print(f"✗ Error loading configuration file: {e}")
        sys.exit(1)


def prepare_config_for_llm(config_df: pd.DataFrame) -> str:
    """
    Prepare configuration data in a format suitable for LLM
    
    Returns:
        Formatted string representation of configuration
    """
    config_text = []
    
    # Get description columns (exclude PAS Field Name, Data Type, Field length)
    description_columns = [col for col in config_df.columns 
                          if 'Description' in col]
    
    config_text.append("CONFIGURATION FILE STRUCTURE:")
    config_text.append("="*80)
    config_text.append(f"\nAvailable instruction columns: {len(description_columns)}")
    for i, col in enumerate(description_columns, 1):
        # Count how many fields have descriptions in this column
        field_count = config_df[col].notna().sum()
        config_text.append(f"{i}. {col} ({field_count} fields with instructions)")
    
    config_text.append("\n" + "="*80)
    config_text.append("FIELD EXTRACTION INSTRUCTIONS:")
    config_text.append("="*80)
    
    # For each PAS field, show all available descriptions
    for idx, row in config_df.iterrows():
        field_name = row['PAS Field Name']
        config_text.append(f"\n[FIELD: {field_name}]")
        
        # Show description from each column
        for col in description_columns:
            description = row.get(col, '')
            if pd.notna(description) and str(description).strip():
                config_text.append(f"  • {col}: {str(description)}")
    
    return "\n".join(config_text)


# =============================================================================
# LLM-BASED INTELLIGENT COLUMN SELECTION
# =============================================================================

def prepare_config_for_llm(config_df: pd.DataFrame) -> str:
    """
    Prepare configuration data with ALL description columns for LLM to analyze and choose from
    
    Args:
        config_df: Configuration DataFrame with all columns
        
    Returns:
        Formatted string representation of complete configuration
    """
    config_text = []
    
    # Get description columns (exclude PAS Field Name, Data Type, Field length)
    description_columns = [col for col in config_df.columns 
                          if 'Description' in col]
    
    config_text.append("CONFIGURATION FILE STRUCTURE:")
    config_text.append("="*80)
    config_text.append(f"\nAvailable instruction columns: {len(description_columns)}")
    for i, col in enumerate(description_columns, 1):
        # Count how many fields have descriptions in this column
        field_count = config_df[col].notna().sum()
        config_text.append(f"{i}. {col} ({field_count} fields with instructions)")
    
    config_text.append("\n" + "="*80)
    config_text.append("FIELD EXTRACTION INSTRUCTIONS:")
    config_text.append("="*80)
    
    # For each PAS field, show all available descriptions
    for idx, row in config_df.iterrows():
        field_name = row['PAS Field Name']
        config_text.append(f"\n[FIELD: {field_name}]")
        
        # Show description from each column
        for col in description_columns:
            description = row.get(col, '')
            if pd.notna(description) and str(description).strip():
                config_text.append(f"  • {col}: {str(description)}")
    
    return "\n".join(config_text)


# =============================================================================
# GPT EXTRACTION - LLM INTELLIGENTLY SELECTS COLUMN AND EXTRACTS FIELDS
# =============================================================================

def extract_fields_with_intelligent_selection(document_text: str, 
                                              config_structure: str,
                                              document_name: str, 
                                              file_extension: str,
                                              pas_fields: List[str],
                                              api_key: str, 
                                              model: str) -> Tuple[Dict[str, str], str]:
    """
    Extract all fields from document using a SINGLE GPT API call
    LLM intelligently analyzes document and selects the most appropriate configuration column
    
    Args:
        document_text: Full text content of the document
        config_structure: Complete configuration with ALL description columns formatted for LLM
        document_name: Name of the document (without extension)
        file_extension: File extension (.pdf, .docx, etc.)
        pas_fields: List of all PAS field names
        api_key: OpenAI API key
        model: Model to use (gpt-4o or gpt-4o-mini)
        
    Returns:
        Tuple of (extracted_fields_dict, selected_column_name)
    """
    
    # Truncate document if too long
    if len(document_text) > MAX_DOCUMENT_CHARS:
        print(f"  ℹ Document truncated from {len(document_text)} to {MAX_DOCUMENT_CHARS} characters")
        document_text = document_text[:MAX_DOCUMENT_CHARS]
    
    # Build the extraction prompt - LLM will analyze and choose column
    prompt = f"""You are an expert document field extraction system with intelligent configuration selection.

TASK OVERVIEW:
You will receive a document and a complete configuration file with multiple instruction columns. Your job is to:
1. Analyze the document name, type, and content to select the MOST APPROPRIATE instruction column
2. Extract all fields using the instructions from that selected column
3. Return the extracted data along with the column you selected

DOCUMENT INFORMATION:
- Document Name: {document_name}
- File Extension: {file_extension}
- Document Type: {"Word Document" if file_extension in ['.docx', '.doc'] else "PDF Document" if file_extension == '.pdf' else "Excel Spreadsheet"}

COLUMN SELECTION GUIDELINES:
1. Analyze the document name and content to determine document type
2. Match to the most appropriate instruction column:
   - If document name contains "PD" AND file is Word (.docx/.doc) → Use "PD (Word Doc) Description"
   - If document name contains "PD" AND file is NOT Word → Use "PD Description"  
   - If document name contains "CAM" → Use "CAM Description"
   - If document name contains "Application" or "Form" → Use "Application Form Description"
   - If document name contains "Legal" → Use "Legal Doc Description"
   - If document name contains "Technical" → Use "Technical Doc Description"
   - If document name contains "Email" + "Subject" → Use "Email Subject Description"
   - If document name contains "Email" + "Body" → Use "Email Body Description"
3. If no clear match, analyze the document content to determine type and select appropriate column
4. Choose the column with the most relevant instructions for this document type

{config_structure}

DOCUMENT CONTENT:
{document_text}

EXTRACTION INSTRUCTIONS:
1. FIRST: Analyze the document name "{document_name}", type "{file_extension}", and content
2. SELECT the most appropriate instruction column based on analysis
3. For each of the {len(pas_fields)} PAS fields, extract the value using instructions from your selected column
4. If a field has no instruction in the selected column, mark as "NO INSTRUCTION"
5. If a field has instruction but value is not found in document, mark as "NOT FOUND"
6. Extract exact values as they appear in the document
7. Do NOT make assumptions or infer values not explicitly stated
8. For numeric values, extract only numbers (remove currency symbols, commas)
9. For dates, maintain format as shown in document

OUTPUT FORMAT:
Return a JSON object with TWO keys:
1. "selected_column": The name of the configuration column you selected (must match exactly)
2. "extracted_fields": An object with field names as keys and extracted values as values

Example format:
{{
  "selected_column": "Application Form Description",
  "extracted_fields": {{
    "Loan_Number": "3000021473",
    "Sourcing_Region": "North Region",
    "Branch_Name": "Main Branch",
    "Field_Without_Instruction": "NO INSTRUCTION",
    "Field_Not_In_Document": "NOT FOUND"
  }}
}}

CRITICAL: 
- Include ALL {len(pas_fields)} PAS fields in extracted_fields, even if "NO INSTRUCTION" or "NOT FOUND"
- The selected_column MUST be one of the column names from the configuration
- Analyze the document thoroughly before selecting the column

Return ONLY the JSON object, no additional text."""

    try:
        client = OpenAI(api_key=api_key)
        
        print(f"  → Making GPT API call with full configuration...")
        print(f"  → LLM will intelligently analyze and select appropriate column...")
        
        response = client.chat.completions.create(
            model=model,
            messages=[
                {
                    "role": "system", 
                    "content": "You are a precise document field extraction assistant with intelligent configuration selection. Always respond with valid JSON."
                },
                {
                    "role": "user", 
                    "content": prompt
                }
            ],
            temperature=0.1,
            response_format={"type": "json_object"}
        )
        
        result_text = response.choices[0].message.content
        result_json = json.loads(result_text)
        
        selected_column = result_json.get("selected_column", "UNKNOWN")
        extracted_fields = result_json.get("extracted_fields", {})
        
        # Ensure all PAS fields are present
        for field in pas_fields:
            if field not in extracted_fields:
                extracted_fields[field] = "ERROR"
        
        print(f"  ✓ LLM selected column: '{selected_column}'")
        print(f"  ✓ Successfully extracted {len(extracted_fields)} fields")
        
        return extracted_fields, selected_column
    
    except json.JSONDecodeError as e:
        print(f"  ✗ JSON parsing error: {e}")
        return {field: "ERROR" for field in pas_fields}, "ERROR"
    
    except Exception as e:
        print(f"  ✗ API error: {e}")
        return {field: "ERROR" for field in pas_fields}, "ERROR"


# =============================================================================
# RESULT MERGING AND OUTPUT
# =============================================================================

def merge_results_to_excel(all_results: Dict[str, Dict[str, str]], 
                          pas_fields: List[str], 
                          output_path: str,
                          column_selections: Dict[str, str]) -> None:
    """
    Merge all extraction results into a single Excel file
    
    Args:
        all_results: Dictionary mapping document names to their extracted fields
        pas_fields: List of all PAS field names
        output_path: Path to save output Excel file
        column_selections: Dictionary mapping document names to selected columns
    """
    results_data = []
    
    for field in pas_fields:
        row = {'PAS Field Name': field}
        for doc_name, extracted_data in all_results.items():
            row[doc_name] = extracted_data.get(field, "NOT PROCESSED")
        results_data.append(row)
    
    results_df = pd.DataFrame(results_data)
    
    # Create metadata sheet with column selections
    metadata_data = []
    for doc_name, selected_col in column_selections.items():
        metadata_data.append({
            'Document Name': doc_name,
            'Selected Configuration Column': selected_col
        })
    metadata_df = pd.DataFrame(metadata_data)
    
    # Save to Excel with multiple sheets
    try:
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            results_df.to_excel(writer, sheet_name='Extracted Fields', index=False)
            metadata_df.to_excel(writer, sheet_name='Column Selections', index=False)
        
        print(f"\n✓ Results saved to: {output_path}")
        print(f"  - Sheet 1: Extracted Fields")
        print(f"  - Sheet 2: Column Selections (shows which config column was used)")
    except Exception as e:
        print(f"\n✗ Error saving results: {e}")


def display_summary(all_results: Dict[str, Dict[str, str]], 
                   pas_fields: List[str],
                   column_selections: Dict[str, str]) -> None:
    """Display summary of extraction results"""
    print("\n" + "="*80)
    print("EXTRACTION SUMMARY")
    print("="*80)
    print(f"Total PAS Fields: {len(pas_fields)}")
    print(f"Total Documents Processed: {len(all_results)}")
    
    print(f"\nColumn Selections (by LLM Intelligent Analysis):")
    for i, (doc_name, selected_col) in enumerate(column_selections.items(), 1):
        print(f"  {i}. {doc_name:30s} → {selected_col}")
    
    # Count extraction statistics
    total_extracted = 0
    total_not_found = 0
    total_no_instruction = 0
    total_errors = 0
    
    for doc_name, extracted_data in all_results.items():
        for value in extracted_data.values():
            if value == "NOT FOUND":
                total_not_found += 1
            elif value == "NO INSTRUCTION":
                total_no_instruction += 1
            elif value == "ERROR":
                total_errors += 1
            elif value != "NOT PROCESSED":
                total_extracted += 1
    
    total_cells = len(pas_fields) * len(all_results)
    
    print(f"\nExtraction Statistics:")
    print(f"  Fields Extracted: {total_extracted} ({total_extracted/total_cells*100:.1f}%)")
    print(f"  Not Found: {total_not_found} ({total_not_found/total_cells*100:.1f}%)")
    print(f"  No Instruction: {total_no_instruction} ({total_no_instruction/total_cells*100:.1f}%)")
    print(f"  Errors: {total_errors} ({total_errors/total_cells*100:.1f}%)")
    print("="*80)


# =============================================================================
# MAIN EXECUTION
# =============================================================================

def main():
    """Main execution function - processes all folders recursively"""
    print("\n" + "="*80)
    print("DOCUMENT FIELD EXTRACTION SYSTEM")
    print("LLM-Based Intelligent Column Selection")
    print("Recursive Folder Processing")
    print("="*80)
    print(f"\nTimestamp: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    # Validate API key
    if not OPENAI_API_KEY or OPENAI_API_KEY == "":
        print("\n✗ ERROR: OpenAI API key not set!")
        print("  Please set the OPENAI_API_KEY environment variable:")
        print("  export OPENAI_API_KEY='your-api-key-here'")
        sys.exit(1)
    
    # Load configuration
    print(f"\n[1/4] Loading Configuration")
    print(f"  Configuration file: {CONFIG_FILE_PATH}")
    
    if not os.path.exists(CONFIG_FILE_PATH):
        print(f"  ✗ Configuration file not found: {CONFIG_FILE_PATH}")
        sys.exit(1)
    
    config_df, pas_fields = load_configuration(CONFIG_FILE_PATH)
    print(f"  ✓ Loaded {len(pas_fields)} PAS fields")
    print(f"  ✓ Available columns: {len(config_df.columns)}")
    
    # Get list of description columns
    description_columns = [col for col in config_df.columns if 'Description' in col]
    print(f"  ✓ Found {len(description_columns)} description columns:")
    for i, col in enumerate(description_columns, 1):
        print(f"    {i}. {col}")
    
    # Prepare complete configuration for LLM
    print(f"\n  Preparing complete configuration for LLM...")
    config_structure = prepare_config_for_llm(config_df)
    print(f"  ✓ Configuration prepared ({len(config_structure)} characters)")
    
    # Discover all folders with documents
    print(f"\n[2/4] Discovering Document Folders")
    print(f"  Base folder: {BASE_DOCUMENTS_FOLDER}")
    
    if not os.path.exists(BASE_DOCUMENTS_FOLDER):
        print(f"  ✗ Base folder not found!")
        print(f"  Creating folder: {BASE_DOCUMENTS_FOLDER}")
        os.makedirs(BASE_DOCUMENTS_FOLDER)
        print(f"  Please place your documents in subfolders and run again.")
        print(f"\n  Example structure:")
        print(f"    {BASE_DOCUMENTS_FOLDER}/")
        print(f"    ├── 3000023987/")
        print(f"    │   └── 20251207_173848/")
        print(f"    │       ├── App_Form.pdf")
        print(f"    │       └── PD_Note.pdf")
        print(f"    └── 3000023988/")
        print(f"        └── 20251207_120000/")
        print(f"            └── Document.pdf")
        sys.exit(1)
    
    supported_extensions = ['.pdf', '.docx', '.doc', '.xlsx', '.xls']
    document_folders = find_document_folders(BASE_DOCUMENTS_FOLDER, supported_extensions)
    
    if not document_folders:
        print(f"  ✗ No folders with documents found in {BASE_DOCUMENTS_FOLDER}")
        print(f"  Supported formats: {', '.join(supported_extensions)}")
        print(f"\n  Please create subfolders and add documents:")
        print(f"    {BASE_DOCUMENTS_FOLDER}/folder1/document.pdf")
        print(f"    {BASE_DOCUMENTS_FOLDER}/folder2/form.docx")
        sys.exit(1)
    
    print(f"  ✓ Found {len(document_folders)} folder(s) with documents:")
    for i, folder in enumerate(document_folders, 1):
        rel_path = os.path.relpath(folder, BASE_DOCUMENTS_FOLDER)
        doc_count = len(get_documents_in_folder(folder, supported_extensions))
        print(f"    {i}. {rel_path} ({doc_count} documents)")
    
    # Process each folder
    print(f"\n[3/4] Processing All Folders")
    print(f"  Using model: {OPENAI_MODEL}")
    print(f"  Output filename: {OUTPUT_FILENAME}")
    print(f"  Strategy: LLM intelligently selects column based on document analysis")
    
    successful_folders = []
    failed_folders = []
    total_documents_processed = 0
    
    for folder_idx, folder_path in enumerate(document_folders, 1):
        rel_path = os.path.relpath(folder_path, BASE_DOCUMENTS_FOLDER)
        print(f"\n{'#'*80}")
        print(f"FOLDER {folder_idx}/{len(document_folders)}: {rel_path}")
        print(f"{'#'*80}")
        
        success, message = process_single_folder(
            folder_path=folder_path,
            config_structure=config_structure,
            pas_fields=pas_fields,
            api_key=OPENAI_API_KEY,
            model=OPENAI_MODEL,
            output_filename=OUTPUT_FILENAME
        )
        
        if success:
            successful_folders.append(rel_path)
            # Count documents in this folder
            docs_in_folder = len(get_documents_in_folder(folder_path, supported_extensions))
            total_documents_processed += docs_in_folder
        else:
            failed_folders.append((rel_path, message))
    
    # Final summary
    print(f"\n[4/4] Overall Summary")
    print(f"\n{'='*80}")
    print(f"PROCESSING COMPLETE!")
    print(f"{'='*80}")
    print(f"\nFolders Statistics:")
    print(f"  Total folders scanned: {len(document_folders)}")
    print(f"  Successfully processed: {len(successful_folders)}")
    print(f"  Failed/Skipped: {len(failed_folders)}")
    print(f"  Total documents processed: {total_documents_processed}")
    
    if successful_folders:
        print(f"\n✓ Successfully Processed Folders:")
        for i, folder in enumerate(successful_folders, 1):
            output_path = os.path.join(BASE_DOCUMENTS_FOLDER, folder, OUTPUT_FILENAME)
            print(f"  {i}. {folder}")
            print(f"     Output: {output_path}")
    
    if failed_folders:
        print(f"\n✗ Failed/Skipped Folders:")
        for i, (folder, reason) in enumerate(failed_folders, 1):
            print(f"  {i}. {folder} - {reason}")
    
    print(f"\nConfiguration:")
    print(f"  Model used: {OPENAI_MODEL}")
    print(f"  API calls made: {total_documents_processed}")
    print(f"  Estimated cost: ~${total_documents_processed * 0.019:.2f}")
    
    print(f"\nNext Steps:")
    print(f"  1. Check output files in each folder: {OUTPUT_FILENAME}")
    print(f"  2. Review Sheet 1 for extracted field values")
    print(f"  3. Review Sheet 2 for column selection decisions")
    print(f"{'='*80}\n")


# =============================================================================
# ENTRY POINT
# =============================================================================

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n\n✗ Process interrupted by user")
        sys.exit(1)
    except Exception as e:
        print(f"\n✗ Unexpected error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
