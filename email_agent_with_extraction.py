#!/usr/bin/env python3
"""
Combined Email Agent and Document Field Extraction System
- Monitors email inbox, downloads attachments, and saves metadata
- After downloading all attachments for a new email, runs field extraction (LLM-based) for that folder only
- Intelligently matches document names with configuration columns (G onwards)

Usage:
    export OPENAI_API_KEY='your-api-key-here'
    python email_agent_with_extraction.py
"""

import os
import sys
import json
import time
import logging
from pathlib import Path
from datetime import datetime
import argparse
import re
import imaplib
import email
from email.header import decode_header
from complete_email_generator import CompleteEmailGenerator
from dotenv import load_dotenv
import random
import string
from openai import AzureOpenAI

# Document processing imports
import pandas as pd
from openpyxl import load_workbook
from docx import Document
from pypdf import PdfReader
import pdfplumber
from openai import OpenAI

load_dotenv()

# =============================================================================
# CONFIGURATION SECTION
# =============================================================================
CONFIG_FILE_PATH = "FieldConfigrationFile.xlsx"  # Configuration Excel file
OUTPUT_FILENAME = "extraction_results.xlsx"      # Output Excel filename (created in each folder)
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
OPENAI_MODEL = "gpt-4o"
MAX_DOCUMENT_CHARS = 15000
SUPPORTED_EXTENSIONS = ['.pdf', '.docx', '.doc', '.xlsx', '.xls', '.txt']

# =============================================================================
# LLM USAGE LOGGER SETUP
# =============================================================================
llm_logger = logging.getLogger('LLMUsage')
llm_logger.setLevel(logging.INFO)
llm_handler = logging.FileHandler('llm_usage.log', encoding='utf-8')
llm_handler.setFormatter(logging.Formatter('%(asctime)s - %(message)s'))
if not llm_logger.hasHandlers():
    llm_logger.addHandler(llm_handler)

# =============================================================================
# DOCUMENT READING FUNCTIONS
# =============================================================================
def read_pdf_document(file_path: str) -> str:
    text_parts = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"\n{'='*60}\n[PAGE {page_num}]\n{'='*60}\n")
                    text_parts.append(page_text)
                tables = page.extract_tables()
                if tables:
                    for table_num, table in enumerate(tables, 1):
                        text_parts.append(f"\n[TABLE {table_num} ON PAGE {page_num}]\n")
                        for row in table:
                            if row:
                                row_text = " | ".join(str(cell) if cell else "" for cell in row)
                                text_parts.append(row_text)
                        text_parts.append("")
        return "\n".join(text_parts)
    except Exception as e:
        print(f"  ⚠ Error reading PDF {file_path}: {e}")
        return ""

def read_word_document(file_path: str) -> str:
    text_parts = []
    try:
        doc = Document(file_path)
        text_parts.append("="*60)
        text_parts.append("[DOCUMENT CONTENT]")
        text_parts.append("="*60)
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        if doc.tables:
            text_parts.append("\n" + "="*60)
            text_parts.append("[TABLES]")
            text_parts.append("="*60)
            for table_num, table in enumerate(doc.tables, 1):
                text_parts.append(f"\n[TABLE {table_num}]")
                for row_num, row in enumerate(table.rows):
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    text_parts.append(row_text)
        return "\n".join(text_parts)
    except Exception as e:
        print(f"  ⚠ Error reading Word document {file_path}: {e}")
        return ""

def read_excel_document(file_path: str) -> str:
    text_parts = []
    try:
        wb = load_workbook(file_path, data_only=True)
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text_parts.append("="*60)
            text_parts.append(f"[SHEET: {sheet_name}]")
            text_parts.append("="*60)
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                if row_text.strip(" |"):
                    text_parts.append(row_text)
        return "\n".join(text_parts)
    except Exception as e:
        print(f"  ⚠ Error reading Excel file {file_path}: {e}")
        return ""

def read_document(filepath):
    ext = Path(filepath).suffix.lower()
    if ext == '.txt':
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    if ext == '.pdf':
        return read_pdf_document(filepath)
    elif ext in ['.docx', '.doc']:
        return read_word_document(filepath)
    elif ext in ['.xlsx', '.xls']:
        return read_excel_document(filepath)
    else:
        print(f"  ⚠ Unsupported file format: {ext}")
        return ""

# =============================================================================
# CONFIGURATION MANAGEMENT WITH INTELLIGENT COLUMN SELECTION
# =============================================================================
def load_configuration(config_path: str):
    """Load the configuration Excel file"""
    try:
        config_df = pd.read_excel(config_path, sheet_name='Sheet1')
        pas_fields = config_df['PAS Field Name'].tolist()
        return config_df, pas_fields
    except Exception as e:
        print(f"✗ Error loading configuration file: {e}")
        sys.exit(1)

def get_description_columns(config_df: pd.DataFrame) -> list:
    """Get all description columns (from column G onwards, index 6+)"""
    # Columns from index 6 onwards are description columns
    description_columns = [col for col in config_df.columns[6:] if 'Description' in col]
    return description_columns

def prepare_config_for_llm(config_df: pd.DataFrame) -> dict:
    """
    Prepare configuration structure for LLM - needed for compatibility with main.py
    """
    description_columns = get_description_columns(config_df)
    config_structure = {}
    
    for col in description_columns:
        config_structure[col] = {
            'fields': config_df['PAS Field Name'].tolist(),
            'descriptions': config_df[col].tolist()
        }
    
    return config_structure

def match_document_to_column(documents: list, description_columns: list, folder_path: str) -> dict:
    """
    Intelligently match all documents to appropriate description columns using LLM (one-time call)
    
    Args:
        documents: List of document filenames with extensions
        description_columns: List of available description columns
        folder_path: Path to the folder where mapping will be saved
    
    Returns:
        Dictionary mapping {filename: matched_column} or {filename: None} if no match
    """
    # Prepare document information for LLM
    document_info = []
    for doc in documents:
        path_obj = Path(doc)
        document_info.append({
            "filename": doc,
            "name": path_obj.stem,
            "extension": path_obj.suffix
        })
    
    # Create few-shot examples for the LLM
    few_shot_examples = """
Examples of correct mappings:

1. Filename: "mail_subject.txt"
   Extension: .txt
   Column: "Email Subject Description"
   Reasoning: Contains "mail_subject" which clearly indicates email subject content

2. Filename: "mail_body.txt"
   Extension: .txt
   Column: "Email Body Description"
   Reasoning: Contains "mail_body" which clearly indicates email body content

3. Filename: "PD Note Kanhaiya.doc"
   Extension: .doc
   Column: "PD (Word Doc) Description"
   Reasoning: Contains "PD" and has .doc/.docx extension, matching Word document description

4. Filename: "Application Form.pdf"
   Extension: .pdf
   Column: "Application Form Description"
   Reasoning: Contains "Application Form" which directly matches the column name

5. Filename: "CAM Report.pdf"
   Extension: .pdf
   Column: "CAM Description"
   Reasoning: Contains "CAM" which indicates Credit Assessment Memo

6. Filename: "Legal Agreement.pdf"
   Extension: .pdf
   Column: "Legal Doc Description"
   Reasoning: Contains "Legal" indicating legal documentation

7. Filename: "PD Discussion.docx"
   Extension: .docx
   Column: "PD (Word Doc) Description"
   Reasoning: Contains "PD" and has Word document extension (.docx)

8. Filename: "Cam - Niranjan Bhosale"
   Extension: .xlsx
   Column: "CAM Description"
   Reasoning: Contains "CAM" and has Excel document extension (.xlsx)

"""

    # Construct the prompt for LLM
    prompt = f"""You are an expert at matching document filenames to appropriate configuration columns for data extraction.

{few_shot_examples}

Now, match ALL the following documents to the most appropriate columns:

Documents to match:
{json.dumps(document_info, indent=2)}

Available columns:
{json.dumps(description_columns, indent=2)}

Matching Rules:
- Match based on keywords in the document name (e.g., "mail_subject" → "Email Subject Description")
- Consider file extensions critically:
  * .txt files named "mail_subject" → "Email Subject Description"
  * .txt files named "mail_body" → "Email Body Description"
  * .doc/.docx files with "PD" → "PD (Word Doc) Description"
  * .pdf files with "Application Form" → "Application Form Description"
- Look for abbreviations (e.g., "CAM", "PD", "App")
- Match column names that explicitly mention file types (e.g., "Word Doc") only to files with matching extensions
- If no clear match exists for a document, use "NONE" as the value

Return a JSON object where each key is the full filename (with extension) and each value is the matched column name or "NONE".

Example output format:
{{
  "mail_subject.txt": "Email Subject Description",
  "PD Note.docx": "PD (Word Doc) Description",
  "random_file.pdf": "NONE"
}}

Return ONLY valid JSON, no explanations."""

    try:
        # client = OpenAI(api_key=OPENAI_API_KEY)
        # response = client.chat.completions.create(
        #     model=OPENAI_MODEL,
        #     messages=[
        #         {
        #             "role": "system", 
        #             "content": "You are a precise assistant that matches filenames to configuration columns. You carefully consider both the document name and file extension. Return only valid JSON."
        #         },
        #         {
        #             "role": "user", 
        #             "content": prompt
        #         }
        #     ],
        #     temperature=0.0,
        #     response_format={"type": "json_object"}
        # )
        client = AzureOpenAI(
            azure_endpoint="https://qc-tspl-dau-mr.openai.azure.com/",
            api_key="DvskuzopcDYytzJygTQiCl1ikUiT8513H8vfpIwVPZPOnfeHCdZ1JQQJ99BEACHYHv6XJ3w3AAABACOGprIt",
            api_version="2025-01-01-preview",
        )
        model="gpt-4o-mini"
        completion = client.chat.completions.create(
                model=model,
                messages=[
                    {
                    "role": "system", 
                    "content": "You are a precise assistant that matches filenames to configuration columns. You carefully consider both the document name and file extension. Return only valid JSON."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                max_tokens=16384,
                temperature=0.0,
                response_format={"type": "json_object"}
            )
        result_text = completion.choices[0].message.content.strip()
        mapping = json.loads(result_text)

        llm_logger.info(json.dumps({
            "model": model,
            "input_tokens": completion.usage.prompt_tokens,
            "output_tokens": completion.usage.completion_tokens,
            "prompt": prompt,
            "response": result_text
        }))
        
        # Validate and clean the mapping
        validated_mapping = {}
        for filename, column in mapping.items():
            if column in description_columns:
                validated_mapping[filename] = column
                #llm_logger.info(f"LLM Mapping: '{filename}' → '{column}'")
            elif column == "NONE":
                validated_mapping[filename] = None
                #llm_logger.info(f"LLM Mapping: '{filename}' → No match found")
            else:
                validated_mapping[filename] = None
                #llm_logger.warning(f"LLM returned invalid column '{column}' for '{filename}'")
        
        # Save mapping to file with timestamp
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        mapping_file = os.path.join(folder_path, f"document_column_mapping_{timestamp}.json")
        with open(mapping_file, 'w', encoding='utf-8') as f:
            json.dump(validated_mapping, f, indent=2)
        
        # Also save human-readable version
        readable_file = os.path.join(folder_path, f"document_column_mapping_{timestamp}.txt")
        with open(readable_file, 'w', encoding='utf-8') as f:
            for filename, column in validated_mapping.items():
                f.write(f"{filename} -> {column if column else 'NO MATCH'}\n")
        
        return validated_mapping
            
    except Exception as e:
        #llm_logger.error(f"Error in LLM batch matching: {e}")
        print(f"  ⚠ Error using LLM for document matching: {e}")
        return {doc: None for doc in documents}

def extract_fields_with_intelligent_selection(
    document_text: str,
    config_df: pd.DataFrame,
    matched_column: str,
    pas_fields: list,
    api_key: str,
    model: str = "gpt-4o",
    document_name: str = None,  # Keep for backward compatibility
    file_extension: str = None   # Keep for backward compatibility
):
    """
    Extract fields from document using the matched description column
    
    Args:
        document_text: The text content of the document
        config_df: Configuration DataFrame
        matched_column: The specific description column to use for extraction
        pas_fields: List of PAS field names
        api_key: OpenAI API key
        model: OpenAI model to use
        document_name: (deprecated, for backward compatibility)
        file_extension: (deprecated, for backward compatibility)
    
    Returns:
        Tuple of (extracted_data dict, selected_column str)
    """
    if not matched_column or matched_column not in config_df.columns:
        print(f"  ⚠ Invalid matched column: {matched_column}")
        return {}, None
    
    # Get the description for the matched column
    descriptions = config_df[matched_column].tolist()
    
    # Build the field extraction prompt
    fields_json = []
    for i, field_name in enumerate(pas_fields):
        description = descriptions[i] if i < len(descriptions) else ""
        if pd.notna(description) and str(description).strip():
            fields_json.append({
                "field_name": field_name,
                "description": str(description)
            })
    
    if not fields_json:
        print(f"  ⚠ No valid field descriptions found in column: {matched_column}")
        return {}, matched_column
    
    # Truncate document if too long
    if len(document_text) > MAX_DOCUMENT_CHARS:
        document_text = document_text[:MAX_DOCUMENT_CHARS] + "\n...[Document truncated]"
    
    prompt = f"""You are a data extraction assistant. Extract the following fields from the document.

Document content:
{document_text}

Fields to extract:
{json.dumps(fields_json, indent=2)}

Instructions:
- For each field, extract the value based on its description
- If a field is not found or not applicable, use null
- Return a JSON object with field names as keys and extracted values as values
- Be precise and only extract what is explicitly mentioned

Return ONLY valid JSON."""

    try:
        client = AzureOpenAI(
            azure_endpoint="https://qc-tspl-dau-mr.openai.azure.com/",
            api_key="DvskuzopcDYytzJygTQiCl1ikUiT8513H8vfpIwVPZPOnfeHCdZ1JQQJ99BEACHYHv6XJ3w3AAABACOGprIt",
            api_version="2025-01-01-preview",
        )

        completion = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert document analysis AI that extracts specific fields from documents based on provided descriptions. Return only valid JSON."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                max_tokens=16384,
                temperature=0.0,
                response_format={"type": "json_object"}
            )
        
        result_text = completion.choices[0].message.content
        extracted_data = json.loads(result_text)
        
        # Log LLM usage
        #llm_logger.info(f"Field extraction using column '{matched_column}': {len(fields_json)} fields processed")
        llm_logger.info(json.dumps({
            "model": model,
            "input_tokens": completion.usage.prompt_tokens,
            "output_tokens": completion.usage.completion_tokens,
            "prompt": prompt,
            "response": result_text
        }))
        return extracted_data, matched_column
        
    except Exception as e:
        print(f"  ✗ Field extraction error: {e}")
        #llm_logger.error(f"Field extraction error with column '{matched_column}': {e}")
        return {}, matched_column

# def merge_results_to_excel(all_results: dict, pas_fields: list, output_path: str, column_selections: dict):
#     """Merge all extraction results into a single Excel file"""
#     try:
#         rows = []
#         for doc_name, extracted_data in all_results.items():
#             row = {"Document": doc_name, "Selected_Column": column_selections.get(doc_name, "Unknown")}
#             for field in pas_fields:
#                 row[field] = extracted_data.get(field, "")
#             rows.append(row)
        
#         df = pd.DataFrame(rows)
#         df.to_excel(output_path, index=False)
#         print(f"  ✓ Results saved to: {output_path}")
#     except Exception as e:
#         print(f"  ✗ Error saving results: {e}")

def merge_results_to_excel(all_results, pas_fields, output_path, column_selections):
    """Merge extraction results into Excel file with custom NA/Not Found logic"""
    try:
        config_df = pd.read_excel(CONFIG_FILE_PATH, sheet_name='Sheet1')
    except Exception as e:
        print(f"  ⚠ Warning: Could not load config file for additional columns: {e}")
        config_df = None

    results_data = []
    doc_columns = list(all_results.keys())
    for field in pas_fields:
        row = {'PAS Field Name': field}
        for doc_name, extracted_data in all_results.items():
            # Determine which column was used for this document
            matched_column = column_selections.get(doc_name)
            description = ""
            if config_df is not None and matched_column in config_df.columns:
                desc_row = config_df[config_df['PAS Field Name'] == field]
                if not desc_row.empty:
                    description = desc_row.iloc[0][matched_column]
            # Custom logic
            if pd.isna(description) or str(description).strip() == "":
                row[doc_name] = "Not Applicable"
            elif field not in extracted_data or extracted_data[field] in [None, "", "null"]:
                row[doc_name] = "Not Found"
            else:
                row[doc_name] = extracted_data[field]
        # Add additional columns from FieldConfiguration file at the end
        if config_df is not None:
            # Find the matching row in config_df for this field
            matching_rows = config_df[config_df['PAS Field Name'] == field]
            if not matching_rows.empty:
                config_row = matching_rows.iloc[0]

                # Compute Final Data for PAS System using First/Second Preference and document values
                def _is_valid_value(value):
                    if value is None:
                        return False
                    if isinstance(value, float) and pd.isna(value):
                        return False
                    text = str(value).strip()
                    if not text:
                        return False
                    if text.upper() in {'NO INSTRUCTION', 'NOT FOUND', 'NOT PROCESSED'}:
                        return False
                    return True

                first_pref = config_row.get('First Preference') if 'First Preference' in config_df.columns else None
                second_pref = config_row.get('Second Preference') if 'Second Preference' in config_df.columns else None

                final_value = None

                # Try First Preference
                if pd.notna(first_pref):
                    pref_col = str(first_pref)
                    if pref_col in doc_columns:
                        candidate = row.get(pref_col)
                        if _is_valid_value(candidate):
                            final_value = candidate

                # Fallback to Second Preference
                if final_value is None and pd.notna(second_pref):
                    pref_col = str(second_pref)
                    if pref_col in doc_columns:
                        candidate = row.get(pref_col)
                        if _is_valid_value(candidate):
                            final_value = candidate

                # Fallback to first valid document value
                if final_value is None:
                    for col in doc_columns:
                        candidate = row.get(col)
                        if _is_valid_value(candidate):
                            final_value = candidate
                            break

                row['Final Data for PAS System'] = final_value if final_value is not None else ""

                # Populate Email Subject / Email Body from extracted values of the corresponding documents
                # based on which instruction column was selected during extraction.
                subject_doc = next(
                    (doc for doc, selected in column_selections.items() if selected == 'Email Subject Description' and doc in doc_columns),
                    None,
                )
                body_doc = next(
                    (doc for doc, selected in column_selections.items() if selected == 'Email Body Description' and doc in doc_columns),
                    None,
                )

                subject_val = row.get(subject_doc) if subject_doc else None
                body_val = row.get(body_doc) if body_doc else None

                row['Email Subject'] = str(subject_val).strip() if _is_valid_value(subject_val) else ""
                row['Email Body'] = str(body_val).strip() if _is_valid_value(body_val) else ""

                # Add selected columns from the configuration file, excluding unwanted metadata/description columns
                excluded_columns = {
                    'Data Type',
                    'Field length',
                    'Primary Source Document',
                    'Secondary Source Document',
                    'CAM Description',
                    'PD Description',
                    'PD (Word Doc) Description',
                    'Application Form Description',
                    'Legal Doc Description',
                    'Technical Doc Description',
                    'Email Subject Description',
                    'Email Body Description',
                }
                for col in config_df.columns:
                    if col == 'PAS Field Name':
                        continue
                    if col in excluded_columns:
                        continue
                    if col == 'Criticality':
                        row['Mismatch Criticality'] = config_row[col] if pd.notna(config_row[col]) else ""
                        continue
                        row[col] = config_row[col] if pd.notna(config_row[col]) else ""
        results_data.append(row)

    results_df = pd.DataFrame(results_data)
    try:
        required_cols = ['PAS Field Name', 'Final Data for PAS System']
        missing_cols = [c for c in required_cols if c not in results_df.columns]
        if missing_cols:
            print(f"\n  ⚠ Warning: Could not write PAS JSON map. Missing columns: {missing_cols}")
        else:
            pas_map_df = results_df[required_cols].copy()
            pas_map_df['PAS Field Name'] = pas_map_df['PAS Field Name'].fillna('').astype(str)
            pas_map_df['Final Data for PAS System'] = pas_map_df['Final Data for PAS System'].fillna('').astype(str)
            pas_map = dict(zip(pas_map_df['PAS Field Name'], pas_map_df['Final Data for PAS System']))

            ts = datetime.now().strftime('%Y%m%d_%H%M%S')
            out_dir = os.path.dirname(output_path) or '.'
            json_path = os.path.join(out_dir, f"pas_field_map_{ts}.json")
            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(pas_map, f, indent=2, ensure_ascii=False)
            print(f"\n✓ PAS JSON map saved to: {json_path}")
    except Exception as e:
        print(f"\n  ⚠ Warning: Could not write PAS JSON map: {e}")
    
    try:
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            results_df.to_excel(writer, sheet_name='Extracted Fields', index=False)
        print(f"\n✓ Results saved to: {output_path}")
    except Exception as e:
        print(f"\n✗ Error saving results: {e}")

# =============================================================================
# EMAIL AGENT CLASS
# =============================================================================
class EmailAgentWithExtraction:
    def __init__(self, config):

        self.email_address = config['email']['address']
        print("*****************EmailAddress:",self.email_address)

        self.password = config['email']['password']
        self.imap_server = config['email']['imap_server']
        self.imap_port = config['email'].get('imap_port', 993)
        self.target_subjects = config['agent']['target_subjects']
        if isinstance(self.target_subjects, str):
            self.target_subjects = [self.target_subjects]
        self.loan_id_pattern = config['agent']['loan_id_pattern']
        self.save_location = Path(config['agent']['save_location'])
        self.check_interval = config['agent'].get('check_interval', 60)
        self.mark_as_read = config['agent'].get('mark_as_read', False)
        self.only_unseen = config['agent'].get('only_unseen', True)
        self.processed_emails = set()
        self.mail = None
        self._setup_logging=config['agent'].get('log_file')
        self.save_location.mkdir(parents=True, exist_ok=True)
        # Extraction config
        self.config_df, self.pas_fields = load_configuration(CONFIG_FILE_PATH)
        
        self.mail = None
        
        
        # Load field extraction configuration
        self.config_df, self.pas_fields = load_configuration(CONFIG_FILE_PATH)
        self.description_columns = get_description_columns(self.config_df)
        
        # Setup logging
        self.logger = logging.getLogger('EmailAgent')
        self.logger.setLevel(logging.INFO)
        handler = logging.FileHandler('email_agent.log', encoding='utf-8')
        handler.setFormatter(logging.Formatter('%(asctime)s - %(message)s'))
        if not self.logger.hasHandlers():
            self.logger.addHandler(handler)
        console = logging.StreamHandler()
        console.setFormatter(logging.Formatter('%(message)s'))
        if not any(isinstance(h, logging.StreamHandler) for h in self.logger.handlers):
            self.logger.addHandler(console)

    #####
    def run_email_generation(self, folder_path, extraction_file):
        """Generate and send summary emails after extraction"""
        try:
            self.logger.info(f"\n  [EMAIL GENERATION] Starting email generation for {extraction_file}")
            print("\n  [EMAIL GENERATION] Starting email generation for {extraction_file}")
            config_file = CONFIG_FILE_PATH
            smtp_config = "config.json"
            api_key = OPENAI_API_KEY

            # No need to check for abhl_imgc.json anymore

            generator = CompleteEmailGenerator(
                extraction_file=extraction_file,
                config_file=config_file,
                api_key=api_key,
                smtp_config=smtp_config
            )

            self.logger.info(f"  [EMAIL GENERATION] Recipients loaded: {len(generator.recipients)} recipient(s)")

            result = generator.generate_and_send_all_emails(folder_path, send_emails=True)

            self.logger.info(f"  [EMAIL GENERATION] ✓ Email generation and sending complete")

        except Exception as e:
            self.logger.error(f"  [EMAIL GENERATION ERROR] {e}")
            import traceback
            self.logger.error(f"Traceback: {traceback.format_exc()}")

    #####

    def connect(self, retry_count=3, retry_delay=5):
        for attempt in range(retry_count):
            try:
                print(f"Attempting to connect to {self.imap_server}:{self.imap_port} (Attempt {attempt + 1}/{retry_count})...")
                self.logger.info(f"Connecting to {self.imap_server}:{self.imap_port}...")
                self.mail = imaplib.IMAP4_SSL(self.imap_server, self.imap_port)
                print(f"Attempting to connect to address and password {self.email_address}:{self.password} (Attempt {attempt + 1}/{retry_count})...")
                self.mail.login(self.email_address, self.password)
                self.logger.info(f"[SUCCESS] Connected to {self.imap_server}")
                return True
            except imaplib.IMAP4.error as e:
                self.logger.error(f"[AUTH FAILED] {e}")
                return False
            except Exception as e:
                self.logger.warning(f"Connection attempt {attempt + 1}/{retry_count} failed: {e}")
                if attempt < retry_count - 1:
                    time.sleep(retry_delay)
        self.logger.error("[FAILED] Could not connect after all retries")
        return False

    def disconnect(self):
        if self.mail:
            try:
                self.mail.logout()
                self.logger.info("[DISCONNECTED] Logged out from email server")
            except:
                pass

    def decode_str(self, s):
        if s is None:
            return ""
        if isinstance(s, bytes):
            s = s.decode('utf-8', errors='ignore')
        parts = decode_header(s)
        decoded_parts = []
        for content, encoding in parts:
            if isinstance(content, bytes):
                content = content.decode(encoding or 'utf-8', errors='ignore')
            decoded_parts.append(content)
        return ''.join(decoded_parts)

    def extract_loan_id(self, subject):
        match = re.search(self.loan_id_pattern, subject)
        return match.group(0) if match else self.generate_random_id()

    def generate_random_id(self):
        return ''.join(random.choices(string.ascii_uppercase + string.digits, k=8))

    def process_email(self, email_id, msg):
        subject = self.decode_str(msg.get('Subject', ''))
        from_addr = self.decode_str(msg.get('From', ''))
        date_str = self.decode_str(msg.get('Date', ''))
        print("Subject:",subject)
        print("From:",from_addr)
        print("Date:",date_str)
        # Check if subject matches target subjects
        # if not any(target in subject for target in self.target_subjects):
        #     return []
        
        self.logger.info(f"\n[NEW EMAIL] Subject: {subject}")
        self.logger.info(f"  From: {from_addr}")
        self.logger.info(f"  Date: {date_str}")
        
        loan_id = self.extract_loan_id(subject)
        print("Loan ID:",loan_id)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        folder_name = f"{loan_id}_{timestamp}"
        save_path = os.path.join(self.save_location, folder_name)
        os.makedirs(save_path, exist_ok=True)
        
        # Save email metadata
        metadata = {
            'subject': subject,
            'from': from_addr,
            'date': date_str,
            'loan_id': loan_id,
            'timestamp': timestamp
        }

        print("Saving email metadata to:",os.path.join(save_path, 'email_metadata.json'))
        with open(os.path.join(save_path, 'email_metadata.json'), 'w', encoding='utf-8') as f:
            json.dump(metadata, f, indent=2)
        print("Saved email metadata.")
        # Save email body
        body_content = self.extract_email_body(msg)
        
        with open(os.path.join(save_path, 'mail_body.txt'), 'w', encoding='utf-8') as f:
            f.write(body_content)
        
        # Save email subject
        with open(os.path.join(save_path, 'mail_subject.txt'), 'w', encoding='utf-8') as f:
            f.write(subject)
        
        # Download attachments
        attachments = self.download_attachments(msg, save_path)
        
        if attachments:
            self.logger.info(f"  [SAVED] {len(attachments)} attachment(s) to: {save_path}")
        
        if self.mark_as_read:
            self.mail.store(email_id, '+FLAGS', '\\Seen')
        
        # Run field extraction on this folder
        self.run_field_extraction(save_path)
        
        return attachments

    def extract_email_body(self, msg):
        body = ""
        if msg.is_multipart():
            for part in msg.walk():
                content_type = part.get_content_type()
                if content_type == "text/plain":
                    payload = part.get_payload(decode=True)
                    if payload:
                        body += payload.decode('utf-8', errors='ignore')
        else:
            payload = msg.get_payload(decode=True)
            if payload:
                body = payload.decode('utf-8', errors='ignore')
        return body

    def download_attachments(self, msg, save_path):
        print("Downloading attachments...")
        attachments = []
        for part in msg.walk():
            print("DEBUG: part content-type:", part.get_content_type())
            print("DEBUG: part content-disposition:", part.get('Content-Disposition'))
            print("DEBUG: part filename:", part.get_filename())
            if part.get_content_maintype() == 'multipart':
                continue
            if part.get('Content-Disposition') is None:
                continue
            
            filename = part.get_filename()
            if filename:
                filename = self.decode_str(filename)
                filepath = os.path.join(save_path, filename)
                
                with open(filepath, 'wb') as f:
                    f.write(part.get_payload(decode=True))
                attachments.append(filepath)
                self.logger.info(f"    [ATTACHMENT] {filename}")
        
        return attachments

    def run_field_extraction(self, folder_path):
        """Run field extraction on all documents in the folder"""
        try:
            self.logger.info(f"\n  [EXTRACTION] Starting field extraction for {folder_path}")
            
            # Get all supported documents in the folder
            documents = [
                f for f in os.listdir(folder_path)
                if Path(f).suffix.lower() in SUPPORTED_EXTENSIONS
            ]
            
            if not documents:
                self.logger.info(f"  [EXTRACTION] No supported documents found in {folder_path}")
                return
            
            self.logger.info(f"  [EXTRACTION] Found {len(documents)} document(s) to process")
            
            # ONE-TIME LLM CALL: Match all documents to columns
            filename_to_column = match_document_to_column(
                documents=documents,
                description_columns=self.description_columns,
                folder_path=folder_path
            )
            
            self.logger.info(f"  [EXTRACTION] Document-to-column mapping completed")
            
            # Extract fields from each document using its matched column
            all_results = {}
            column_selections = {}
            
            for doc_file in documents:
                matched_column = filename_to_column.get(doc_file)
                if not matched_column:
                    self.logger.warning(f"    [EXTRACTION] Document '{doc_file}' has no matching column. Skipping.")
                    continue
                
                doc_path = os.path.join(folder_path, doc_file)
                doc_name = Path(doc_file).stem
                
                self.logger.info(f"    [EXTRACTION] Processing: {doc_file} (using column: {matched_column})")
                
                # Read document
                document_text = read_document(doc_path)
                if not document_text or len(document_text) < 50:
                    self.logger.info(f"    [EXTRACTION] Skipping {doc_file} (not enough content)")
                    continue
                
                # Extract fields using the matched column
                extracted_data, _ = extract_fields_with_intelligent_selection(
                    document_text=document_text,
                    config_df=self.config_df,
                    matched_column=matched_column,
                    pas_fields=self.pas_fields,
                    api_key=OPENAI_API_KEY,
                    model=OPENAI_MODEL
                )
                
                if extracted_data:
                    all_results[doc_name] = extracted_data
                    column_selections[doc_name] = matched_column
                    self.logger.info(f"    [EXTRACTION] ✓ Extracted {len(extracted_data)} fields using column: {matched_column}")
            
            # Save results
            if all_results:
                print(all_results)
                output_path = os.path.join(folder_path, OUTPUT_FILENAME)
                merge_results_to_excel(all_results, self.pas_fields, output_path, column_selections)
                self.logger.info(f"  [EXTRACTION] Results saved to {output_path}")
                self.run_email_generation(folder_path, output_path)  ##Added by AP
            else:
                self.logger.info(f"  [EXTRACTION] No extraction results to save for {folder_path}")
                
        except Exception as e:
            self.logger.error(f"  [EXTRACTION ERROR] {e}")
            import traceback
            self.logger.error(f"Traceback: {traceback.format_exc()}")

    def check_emails(self):
        all_attachments = []
        print("Inside check_emails")
        try:
            try:
                self.mail.select('inbox')
            except Exception:
                self.logger.warning("Connection lost, reconnecting...")
                if not self.connect():
                    return all_attachments
                self.mail.select('inbox')
            search_criteria = 'UNSEEN' if self.only_unseen else 'ALL'
            status, messages = self.mail.search(None, search_criteria)
            print("Messages:",len(messages[0].split()))
            print("Status:", status)
            if status != 'OK':
                self.logger.error(f"Search failed: {status}")
                return all_attachments
            email_ids = messages[0].split()
            print("Email IDs:",email_ids)
            for email_id in email_ids:
                if email_id in self.processed_emails:
                    continue
                try:
                    status, msg_data = self.mail.fetch(email_id, '(RFC822)')
                    if status != 'OK':
                        continue
                    for response_part in msg_data:
                        print("Response Part:")
                        if isinstance(response_part, tuple):
                            msg = email.message_from_bytes(response_part[1])
                            #print("Message:",msg)
                            attachments = self.process_email(email_id, msg)
                            print("Attachments:",attachments)
                            all_attachments.extend(attachments)
                    self.processed_emails.add(email_id)
                except Exception as e:
                    self.logger.error(f"[ERROR] Fetching email {email_id.decode('utf-8')}: {e}")
                    continue
        except Exception as e:
            self.logger.error(f"[ERROR] Checking emails: {e}")
        return all_attachments

    def run(self):
        print("Inside run")
        self.logger.info("\n" + "=" * 60)
        self.logger.info("EMAIL AGENT WITH FIELD EXTRACTION STARTED")
        self.logger.info("=" * 60)
        self.logger.info(f"Email: {self.email_address}")
        self.logger.info(f"Server: {self.imap_server}:{self.imap_port}")
        self.logger.info(f"Watching for subjects: {', '.join(self.target_subjects)}")
        self.logger.info(f"Loan ID Pattern: {self.loan_id_pattern}")
        self.logger.info(f"Root save location: {self.save_location}")
        self.logger.info(f"Check interval: {self.check_interval}s")
        self.logger.info(f"Only unseen emails: {self.only_unseen}")
        self.logger.info(f"Mark as read: {self.mark_as_read}")
        self.logger.info("=" * 60 + "\n")
        print("Inside run___ ")
        if not self.connect():
            self.logger.error("[FAILED] Could not start agent - connection failed")
            print("Could not start agent - connection failed")
            return
        
        try:
            while True:
                print("Checking emails...")
                self.logger.info(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Checking emails...")
                attachments = self.check_emails()
                
                if attachments:
                    self.logger.info(f"[SUCCESS] Downloaded {len(attachments)} attachment(s) into {len(set(Path(f).parent for f in attachments))} folder(s)")
                else:
                    self.logger.info("  No new attachments")
                
                time.sleep(self.check_interval)
                
        except KeyboardInterrupt:
            self.logger.info("\n\n[STOPPED] Agent stopped by user")
        except Exception as e:
            self.logger.error(f"\n\n[CRASH] Agent crashed: {e}")
            import traceback
            self.logger.error(f"Traceback: {traceback.format_exc()}")
        finally:
            self.disconnect()

# =============================================================================
# CONFIG LOADING AND MAIN
# =============================================================================
def load_config(config_file='config.json'):
    with open(config_file, 'r', encoding='utf-8') as f:
        return json.load(f)

def main():
    parser = argparse.ArgumentParser(description='Email Agent with Field Extraction')
    parser.add_argument('-c', '--config', default='config.json', help='Configuration file (default: config.json)')
    args = parser.parse_args()
    
    if not os.path.exists(args.config):
        print(f"[ERROR] Config file not found: {args.config}")
        return
    
    if not os.path.exists(CONFIG_FILE_PATH):
        print(f"[ERROR] Field extraction config not found: {CONFIG_FILE_PATH}")
        return
    
    # if not OPENAI_API_KEY:
    #     print(f"[ERROR] OPENAI_API_KEY not set in environment.")
    #     return
    print("Loading config...",args.config)
    config = load_config(args.config)
    agent = EmailAgentWithExtraction(config)
    agent.run()

if __name__ == "__main__":
    main()
